"""
Macro ramp report generation: calls Claude Opus 4.6 with extended thinking + web search.
Exports report as Word doc and PDF.
"""

import os
import re
import anthropic
from datetime import datetime
from pathlib import Path

from gather_macro import web_search

SEARCH_TOOL = {
    "name": "web_search",
    "description": (
        "Search the web for current macro data, central bank policy, positioning, "
        "COT reports, economic releases, analyst forecasts, or geopolitical context. "
        "Use liberally to fill gaps in the provided data."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "The search query"}
        },
        "required": ["query"],
    },
}

FRAMEWORK = """\
You follow this research framework for every macro asset ramp report.

PART 1 — QUICK PASS/FAIL
What kills this trade immediately?
- Is the macro regime directly opposed to this direction?
- Is positioning already at extremes in the direction you want to trade?
- Is the narrative fully priced in with no remaining catalyst?
- Is there an imminent event that invalidates the thesis?
Verdict: Pass / Pass with caveats / Fail — be direct.

PART 2 — EXECUTIVE SUMMARY
- Current regime + asset momentum (1 sentence)
- Key driver and conviction level (1 sentence)
- Trade thesis + expected move + time frame (1 sentence)
- Conclusion: Long / Short / Watch / Avoid

PART 3 — ASSET PROFILE
- What is this asset and how does it trade? (spot, futures, ETF proxies, trading hours)
- Key market participants: commercials, speculators, central banks, real money, retail
- Typical daily/weekly volatility, ATR as % of price, normal move sizes
- Key correlations: list 5+ assets it moves with/against and directionality
- Carry/roll characteristics: cost of carry, futures curve structure if applicable

PART 4 — FUNDAMENTAL DRIVERS
Present this section based on the asset class:

FOR COMMODITIES:
- Supply: production levels, OPEC/producer cartel decisions, supply disruptions, inventory levels (EIA/LME/CFTC stocks) vs 5-year average
- Demand: key demand drivers, geographic concentration (China share), demand trend
- Cost curve: marginal cost of production, where does the cost curve support price?
- Futures curve: contango or backwardation? Implied carry cost/roll yield
- USD sensitivity: historical correlation with DXY, current reading
- Seasonal patterns: typical seasonal move for this time of year

FOR RATES:
- Central bank: current policy rate, last decision, forward guidance, next meeting date, market-implied rate path (show OIS-implied path for next 12 months)
- Inflation: current CPI/PCE level and trend, 5Y5Y breakeven, 2Y2Y breakeven, real yield level vs history
- Growth: current GDP run rate, ISM PMIs, employment trend, leading indicators
- Term premium: current ACM term premium or equivalent, vs history
- Supply/demand: Treasury issuance calendar, Fed QT pace, foreign demand trend (Fed custody holdings), TIC data
- Fiscal: deficit trajectory, debt/GDP, upcoming auctions

FOR FX:
- Rate differential: current short-rate differential (e.g. Fed Funds vs ECB rate), 2Y swap spread, direction of travel
- Real rate differential: inflation-adjusted yield spread vs history
- Growth differential: relative GDP/PMI momentum, which economy is accelerating?
- Current account: trade balance trend, current account surplus/deficit as % GDP for each country
- Central bank divergence: which CB is more hawkish/dovish? Next meeting dates for both
- Positioning: latest IMM speculative net positioning vs 1Y range (z-score)
- PPP fair value: OECD PPP or Big Mac index — % over/undervalued at current spot

FOR EQUITY INDEXES:
- Earnings: current NTM EPS consensus, revision trend (breadth of upgrades vs downgrades), earnings beat rate
- Valuation: forward P/E vs 10Y average, Shiller CAPE, equity risk premium (earnings yield minus real 10Y yield)
- Liquidity: Fed/ECB/BOJ balance sheet direction, M2 growth rate, Goldman/Bloomberg financial conditions index
- Risk appetite: VIX level + trend, HY credit spreads, IG spreads, EM capital flows
- Sector composition: top 3 sectors by weight, top 10 names concentration %, recent sector rotation
- Buyback/dividend yield as price support

PART 5 — MACRO REGIME MAPPING
Map the current environment:
- Inflation/Growth quadrant: Inflationary Boom / Inflationary Bust / Disinflationary Boom / Disinflationary Bust
  → How has this asset class performed historically in this quadrant? Give specific return data.
- Monetary policy cycle: Early tightening / Peak tightening / Early easing / Deep easing
  → How does this asset typically perform at this stage?
- Current narrative: what is the market's dominant thesis? Is it sustainable or approaching exhaustion?
- Regime change risk: what macro data would trigger a regime shift?

PART 6 — POSITIONING & SENTIMENT
- COT report (CFTC): net speculative positioning — current absolute level, vs 1Y range, z-score, direction of recent change
- ETF flows: recent weekly/monthly flows into relevant ETFs (direction, magnitude, trend)
- Options market: current implied vol level, IV rank/percentile vs 1Y, put/call ratio, skew (for FX: risk reversals)
- Analyst/strategist consensus: median target, distribution of forecasts, recent upgrades/downgrades
- Retail sentiment indicators (AAII, CNN Fear & Greed if applicable)
- Crowding assessment: is this a consensus trade? What is the exit risk if wrong?

PART 7 — TECHNICALS
- Current price vs 14d/50d/200d MAs: above/below each, % gap, MA alignment (bullish/bearish/mixed)
- RSI(14): current level, trend, any divergences with price
- Key support levels: S1, S2, S3 with specific prices and reasoning (prior highs/lows, MAs, round numbers, Fibonacci)
- Key resistance levels: R1, R2, R3 with specific prices and reasoning
- Volume/open interest trend: is the move confirmed or suspect?
- Weekly/monthly chart context: is the longer-term trend supportive of the trade?
- Historical analogue: identify a similar price pattern or macro setup from history. What happened? How long did it take?

PART 8 — CROSS-ASSET CONTEXT
- Risk-on / risk-off classification for this asset
- Rolling 30-day correlations to: USD (DXY), S&P 500, 10Y yields, gold, crude oil (use whichever are most relevant)
- Any correlation breakdown vs historical norm? Flag explicitly as a potential signal
- How does this asset fit in a macro portfolio? (diversifier, hedge, risk asset, carry trade)
- Current cross-asset signals that confirm or contradict the thesis

PART 9 — VALUATION / FAIR VALUE
Build a fair value estimate:
- FX: OECD PPP model, BEER model (real effective exchange rate), % over/undervalued, historical mean reversion time
- Rates: Taylor rule implied rate, neutral real rate (r*) estimate, fair yield based on inflation breakevens + term premium
- Commodities: cost of production / marginal cost floor, supply/demand equilibrium price model, historical inflation-adjusted price percentile
- Equity indexes: forward P/E vs historical average + premium/discount, Shiller CAPE implied 10Y return, ERP model (earnings yield vs bond yield)
Show current price vs fair value estimate and the % deviation.

PART 10 — RISKS
What would force a position reversal?
- Bear case: construct the full opposing thesis — don't strawman it
- Key data/events that could invalidate within 2-4 weeks
- Geopolitical tail risks specific to this asset
- Positioning unwind risk: what happens if consensus is wrong?
- Liquidity risk: can you hold through a 2-sigma adverse move? What is the margin/funding cost?
- Correlation breakdown: what if this asset stops behaving as expected?

PART 11 — CATALYST & TRADE SETUP
- Primary catalyst: what drives the move and specific timing
- Key upcoming events calendar (next 60 days): CB meetings, CPI/NFP/GDP prints, OPEC, earnings seasons, elections, Treasury auctions — with dates
- Entry: specific level and timing rationale (do not chase)
- Stop: specific level with reasoning (do NOT place at obvious round numbers or recent extremes — use ATR to set width)
- Targets: T1 (conservative), T2 (base case), T3 (extended) with specific levels and reasoning
- Time frame: expected hold period (days / weeks / months)
- R:R ratio: state explicitly. Minimum 2:1, prefer 3:1+
- Position sizing: rate conviction 1-10, assign 1-4% portfolio risk accordingly
- Hedge: what offsetting position reduces tail risk without killing the trade?
- Options angle: check if implied vol is cheap or expensive (IV rank). Is there a risk reversal or spread structure that improves R:R?
- DCA plan if scaling in
"""


def generate_report(gathered: dict) -> str:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    display_name = gathered.get("display_name", gathered.get("asset_input", "unknown"))
    asset_type   = gathered.get("asset_type", "unknown")

    prompt = f"""You are a sophisticated macro trader and global macro analyst. \
Produce a comprehensive macro ramp report for **{display_name}** ({asset_type.upper()}) \
using the framework below.

{FRAMEWORK}

---

## DATA PROVIDED

### Price Data & Technicals
{gathered.get('price_data') or 'Not available'}

### Macro News & Context
{gathered.get('news') or 'Not available'}

### User-Uploaded Research Files
{gathered.get('uploaded_files') or 'None uploaded'}

### News Digest (Telegram — last 30 days)
{gathered.get('telegram_digest') or 'Not available'}

---

## INSTRUCTIONS

**UPLOADED FILES ARE YOUR PRIMARY SOURCE. READ THEM FIRST, BEFORE ANYTHING ELSE.**
The user has uploaded research files (broker reports, models, positioning data). For every section:
- Extract every number, target, positioning figure, and trade idea from the uploaded files FIRST
- Use uploaded files as the authoritative source — they override yfinance/web data where they conflict
- If a broker report has a specific price target, entry, stop, or thesis: quote it directly and build on it
- If a model is uploaded: extract all key figures (fair value, returns, assumptions) and reference them explicitly
- Do NOT leave any section as "unknown" or "unavailable" if the uploaded files contain the data
- Quote specific page/section references from the uploaded files where helpful

Use the web_search tool liberally to:
- Get latest COT/CFTC positioning data for this asset
- Find current central bank statements, meeting minutes, forward guidance
- Check latest economic data releases (CPI, NFP, PMI, GDP) and surprises
- Find analyst and strategist price targets and recent changes
- Get current futures curve / forward rates / breakeven inflation
- Search for historical analogues and comparable market setups
- Check current ETF flow data and options market positioning

Be specific — use exact numbers, dates, levels, and names. Never be vague.
Where data is genuinely unavailable after searching, say so explicitly with a note on what to check manually.

**CITE YOUR SOURCES THROUGHOUT.**
Every factual claim, number, target, or data point must have a source tag immediately after it. Use these formats:
- Uploaded file: [Source: uploaded broker report] or [Source: uploaded model]
- Web search: [Source: Bloomberg, Mar 2025] or [Source: Reuters] or [Source: web search]
- yfinance/price data: [Source: yfinance]
- COT/CFTC data: [Source: CFTC COT report]
- Central bank: [Source: Fed statement] or [Source: ECB minutes]
- News Digest Telegram: [Source: News Digest]
If a figure appears in both an uploaded file and a web result, cite the uploaded file.
Do NOT write entire paragraphs without any source tags.

Start with Part 1 (Quick Pass/Fail). If this is an obvious fail, say so clearly and briefly.
Then complete all 11 parts of the framework.
"""

    messages = [{"role": "user", "content": prompt}]

    print(f"  Calling Claude Opus with extended thinking...")

    while True:
        response = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=16000,
            thinking={"type": "adaptive"},
            tools=[SEARCH_TOOL],
            messages=messages,
        )
        if response.stop_reason != "tool_use":
            break

        tool_blocks = [b for b in response.content if b.type == "tool_use"]
        print(f"  [web search x{len(tool_blocks)}] {', '.join(b.input['query'] for b in tool_blocks)}")

        def _search(block):
            return {
                "type": "tool_result",
                "tool_use_id": block.id,
                "content": web_search(block.input["query"]),
            }

        from concurrent.futures import ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=len(tool_blocks)) as pool:
            tool_results = list(pool.map(_search, tool_blocks))

        messages.append({"role": "assistant", "content": response.content})
        messages.append({"role": "user", "content": tool_results})

    report_text = ""
    for block in response.content:
        if block.type == "text":
            report_text = block.text.strip()

    return report_text


def export_docx(asset_input: str, display_name: str, report_text: str, output_dir: str) -> str:
    from docx import Document
    from docx.shared import Pt

    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_name = re.sub(r'[\\/:*?"<>|]', "", display_name).strip().replace(" ", "_")
    filename = f"{asset_input.upper()}_{safe_name}_{date_str}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_runs(para, text):
        for part in re.split(r"(\*\*[^*]+\*\*|__[^_]+__)", text):
            if (part.startswith("**") and part.endswith("**")) or \
               (part.startswith("__") and part.endswith("__")):
                para.add_run(part[2:-2]).bold = True
            else:
                para.add_run(part)

    def add_table(table_lines):
        parsed = []
        for tl in table_lines:
            if re.match(r"^\s*\|[-| :]+\|\s*$", tl):
                continue
            cells = [c.strip() for c in tl.strip("|").split("|")]
            parsed.append(cells)
        if not parsed:
            return
        ncols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=ncols)
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(parsed):
            for c_idx in range(ncols):
                text = row[c_idx].replace("**", "").replace("__", "") if c_idx < len(row) else ""
                cell = tbl.cell(r_idx, c_idx)
                cell.text = ""
                run = cell.paragraphs[0].add_run(text)
                if r_idx == 0:
                    run.bold = True
        doc.add_paragraph()

    lines = report_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(block)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif not stripped or stripped in ("---", "==="):
            doc.add_paragraph()
        elif re.match(r"^[-*] ", stripped):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\. ", "", stripped))
        else:
            add_runs(doc.add_paragraph(), stripped)

        i += 1

    doc.save(filepath)
    print(f"  Docx saved: {filepath}")
    return filepath


def export_pdf(docx_path: str) -> str:
    import platform, subprocess
    pdf_path = docx_path.replace(".docx", ".pdf")
    if platform.system() == "Windows":
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    else:
        output_dir = str(Path(docx_path).parent)
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            check=True, capture_output=True,
        )
    print(f"  PDF saved: {pdf_path}")
    return pdf_path
