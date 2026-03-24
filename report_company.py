"""
Report generation: calls Claude Opus 4.6 with extended thinking + web search tool.
Exports report as PDF using fpdf2.
"""

import os
import anthropic
from datetime import datetime
from pathlib import Path

from gather_company import web_search

SEARCH_TOOL = {
    "name": "web_search",
    "description": (
        "Search the web for current information about companies, industries, competitors, "
        "management, SEC filings, earnings calls, analyst reports, or macro context. "
        "Use liberally to fill gaps in the provided data."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "query": {"type": "string", "description": "The search query"}
        },
        "required": ["query"],
    },
}

FRAMEWORK = """\
You follow this research framework for every company ramp report:

PART 1 — QUICK PASS/FAIL
Answer: What would kill this idea? Try to rule it out quickly.
Check: Management track record on M&A/turnarounds, earnings call clarity (do they dodge questions?),
Glassdoor culture signals, intercompany transaction risks, opposing thesis not yet bottomed,
cannibalization of own products, new entrant risk ("your margin is my opportunity").
Verdict: Pass / Pass with caveats / Fail — explain briefly.

PART 2 — EXECUTIVE SUMMARY
- Business quality + current market sentiment (1 sentence)
- Valuation vs peers (1 sentence)
- Catalyst + expected upside over time frame (1 sentence)
- Conclusion: Buy / Watch / Avoid

PART 3 — COMPANY DEEP-DIVE
- Value proposition and what the company does simply
- Revenue/business model: how do they make money?
- Product segments and geographic distribution
- End customer segments, customer concentration, GTM/sales process
- Supply chain: key suppliers, contracts, production capabilities
- Management: background, compensation, ownership stake, M&A history, track record
- Share price history: what has historically moved the stock?
- Shareholding structure (institutional, insider, short interest)

PART 4 — FINANCIALS

4A. REVENUE BREAKDOWN
- If the company reports multiple segments: show each segment's revenue, % of total, YoY growth rate, and CAGR over the available history. Use a table.
- If the company is single-segment: note this explicitly ("Single reportable segment — no breakdown available") and provide total revenue history instead.
- Include geographic breakdown if disclosed.
- Use the SEC filing data provided (3 years of history). Fill forward estimates from web search / guidance.

4B. GAAP AND NON-GAAP INCOME STATEMENTS

Present four tables in this exact order:

TABLE 1 — GAAP Income Statement ($M)
Rows: Revenue | COGS | Gross Profit | R&D | Sales & Marketing | G&A | Total OpEx (below GP) | EBIT | Interest Income | Interest Expense | Other Income/(Expense) | EBT | Tax | Net Income
Columns: each fiscal year available + NTM/forward estimate.
Derive COGS from Revenue × (1 − Gross Margin%) if not directly stated. Show losses in parentheses.

TABLE 2 — GAAP Common-Sized (% of Revenue)
Same rows as Table 1. Every line expressed as % of revenue.

TABLE 3 — Non-GAAP Income Statement ($M)
Rows: Revenue | Adj. COGS | Adj. Gross Profit | Adj. R&D | Adj. S&M | Adj. G&A | Adj. EBITDA | D&A | Adj. EBIT | Interest & Other | Adj. EBT | Adj. Net Income | Adj. EPS | Adj. Diluted Shares

TABLE 4 — Non-GAAP Common-Sized (% of Revenue)
Same rows as Table 3. Every line expressed as % of revenue.

LINE ITEM COMPOSITION (written explanation below the four tables — do NOT indent into the tables):
After the four tables, write a section called "Line Item Breakdown" that explains what each line is made up of, based on what the company actually discloses in its filings and earnings releases. Cover every line item. Format as:

**Revenue:** [what product/service lines or segments make up revenue, relative sizes]
**COGS:** [what is inside cost of revenue — e.g. hosting/infrastructure, personnel costs, amortization of capitalized software, support costs — with approximate $ or % split where disclosed]
**R&D:** [personnel, SBC, compute/cloud costs for R&D, any other disclosed items]
**Sales & Marketing:** [personnel, commissions, SBC, advertising/demand gen, other]
**G&A:** [personnel, SBC, legal/audit/insurance, other]
**SBC total:** [total stock-based comp and how it is split across COGS/R&D/S&M/G&A]
**D&A:** [depreciation split by asset type if disclosed — GPU/server depreciation, real estate, software; amortization of acquired intangibles]
**Interest Income:** [cash/money market yield on cash balance]
**Interest Expense:** [debt outstanding, rate, any convertible notes]
**Tax:** [effective tax rate GAAP vs non-GAAP, deferred tax assets/NOL position]

Then add commentary:
- COGS composition: what are the biggest cost drivers — fixed vs variable, and how do they scale with revenue?
- GAAP vs non-GAAP gap: what is excluded (SBC, D&A, restructuring)? Aggressive or reasonable?
- Gross margin expansion: which specific sub-components are driving improvement?
- Opex leverage: which line items are scaling efficiently vs growing faster than revenue?
- Breakeven path: when does each major cost line cross a leverage threshold?

MANAGEMENT EXPLANATION OF CHANGES (written section immediately after Line Item Breakdown):
Pull directly from earnings call transcripts, press releases, and MD&A sections. For each major YoY change in revenue and margins, quote or closely paraphrase what management said was the reason. Structure as:

**Revenue growth/decline — management explanation:** [exact reason given: new products, pricing, volume, geography, churn, macro, FX — with specific numbers cited by management]
**Gross margin change — management explanation:** [what management said drove the expansion or compression: mix shift, pricing power, infra efficiency, GPU depreciation, headcount, hyperscaler discounts, etc.]
**R&D change — management explanation:** [headcount additions, new product investments, any commentary on R&D efficiency]
**S&M change — management explanation:** [go-to-market changes, sales headcount, marketing spend, CAC trends]
**G&A change — management explanation:** [any scale efficiencies, one-time items, compliance costs]
**Forward guidance rationale:** [what management said about why they guided the way they did — conservatism, visibility, known headwinds/tailwinds]

Flag any discrepancy between what management said and what the numbers show (e.g. management claims gross margin improvement due to efficiency but COGS grew faster than revenue).

4C. CREDIT METRICS
Show BOTH sets of metrics — financial debt only AND lease-adjusted:
  Financial debt metrics: Net Financial Debt / EBITDA, Interest Coverage (EBIT / Interest), Total Financial Debt / Equity
  Lease-adjusted metrics: (Net Debt + Operating Lease Obligations) / EBITDA, Fixed Charge Coverage
If zero financial debt: explicitly state this and flag as a balance sheet positive, but still show lease-adjusted leverage.
Also show: Cash + equivalents, Revolver capacity (drawn vs undrawn), FCF conversion (FCF / Net Income), Capex / Revenue trend.

4D. CASH FLOW
- Operating cash flow, capex, free cash flow (last 3 years + NTM estimate)
- Capex intensity (capex / revenue) and whether it is declining or rising
- NWC dynamics: receivables days, inventory days, payables days if applicable
- Reinvestment rate and return on invested capital (ROIC)

4E. LTM AND NTM MULTIPLES
EV/Revenue, EV/EBITDA, EV/EBIT, P/E, P/FCF — LTM and NTM side by side.

4F. 10-YEAR DCF
Build a full 10-year DCF. Show all assumptions explicitly in a table:
  - Revenue growth rate by year (use guidance for Y1-2, taper to terminal for Y3-10)
  - EBITDA margin by year
  - D&A, Capex, Change in NWC assumptions
  - Tax rate
  - WACC: calculate explicitly — risk-free rate (current 10Y UST), equity risk premium (use 5.5%), beta (from data provided), cost of debt, target capital structure → show the WACC build
  - Terminal value: calculate using BOTH Gordon Growth Model (terminal growth rate = 2.5%) AND Exit Multiple (EV/EBITDA). Show both implied share prices.
  - Sensitivity table 1: implied share price vs Revenue CAGR (rows) × EBITDA margin (columns) — 3×3 grid
  - Sensitivity table 2: implied share price vs WACC (rows) × Exit Multiple / Terminal Growth (columns) — 3×3 grid
  - State clearly what growth and margin assumptions are currently priced in at the current share price.

4G. MANAGEMENT GUIDANCE AND KEY SWING FACTORS
- Full guidance table (revenue, EPS, margins, capex) with actuals vs consensus
- Top 3 swing factors that could cause the stock to beat or miss

PART 5 — INDUSTRY & COMPETITIVE POSITION
- Industry trends and TAM (size + historical/forward growth)
- Competitive landscape: segment leaders, challengers, disruptors
- Company's differentiation vs competitors (qualitative + metrics)
- Market share and trend
- Regulatory environment

PART 6 — MACRO REGIME CONTEXT
Map the company into the current macro regime:
Inflation vs Growth quadrant: Inflationary Boom / Inflationary Bust / Disinflationary Boom / Disinflationary Bust
  → Playbook implications for this asset class
Monetary Policy vs Growth quadrant: Slowing+Tightening / Slowing+Easing / Strong+Tightening / Strong+Easing
  → Playbook implications
Current narrative: Is the theme driving this stock sustainable? What kills the narrative? What is next?
Sector tailwinds/headwinds from macro backdrop.

PART 7 — VALUATION
- Screen comparable companies (same sector, geography, business model, size)
- Explain WHY multiples differ — growth-adjusted view (PEG, growth-adjusted EV/EBITDA)
- Low/high multiples are appropriate at the right growth rate; don't exclude outliers blindly
- DCF sanity check: what growth and margin assumptions are currently priced in?
- Historical valuation range for the stock

PART 8 — TECHNICALS & POSITIONING
- Price vs 14d / 50d / 200d moving averages — where is it in the range?
- RSI and momentum signals
- Key support / resistance levels from recent price action
- Historical analogue: find a similar setup in market history and what happened
- Positioning: short interest trend, put/call ratio, COT if applicable
- Sentiment: analyst consensus direction, recent upgrades/downgrades
- Rule: positioning/sentiment are trend-following signals; fade only if extreme or new event will unwind them

PART 9 — RISKS
Answer: Why would someone sell this position to me?
- Opposing thesis: has the bearish case bottomed out? Why now?
- Cannibalization / intercompany transaction risks
- Management fog: do they understand their own business? Do they follow through on guidance?
- New entrant risk: is the margin too attractive to ignore?
- Macro tail risks specific to this company (FX, rate sensitivity, credit, commodity exposure)
- Structural change risk: is this a value trap or secular decline story?

PART 10 — CATALYST & TRADE SETUP
- Primary catalyst: what drives the narrative shift and when?
- Key upcoming events: earnings date, product launches, regulatory decisions, macro data
- Entry thesis: why now specifically?
- Technical entry: do not place stops at obvious levels; use recent volatility to set width
- Position sizing: rate conviction 1-10, assign 1-4% portfolio risk accordingly
  (A 5/10 → 50-80% upside potential; a 4% position = comfortable buying 20% lower)
- DCA approach if sizing up slowly
- Options angle: if bimodal event, check if market is pricing normal distribution (mispricing opportunity)
  Long-dated low-IV options may be underpriced for structural themes
- Correlations to check: existing portfolio exposure by geography, sector, asset class
"""


def generate_report(ticker: str, gathered: dict) -> str:
    client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

    company_name = gathered.get("company_name", ticker)

    prompt = f"""You are a sophisticated macro trader and equity analyst. Produce a comprehensive company ramp report for **{ticker} ({company_name})** using the framework below.

{FRAMEWORK}

---

## DATA PROVIDED

### Yahoo Finance Data
{gathered.get('yfinance') or 'Not available'}

### Web News & Analyst Coverage
{gathered.get('news') or 'Not available'}

### SEC EDGAR Filings (8-K Press Release, 10-Q, 10-K)
{gathered.get('sec') or 'Not available'}

### Earnings Call Transcript
{gathered.get('transcript') or 'Not available'}

### User-Uploaded Research Files
{gathered.get('uploaded_files') or 'None uploaded'}

### Stock Digest (Telegram — last 1 year)
{gathered.get('telegram_digest') or 'Not available'}

---

## INSTRUCTIONS

**UPLOADED FILES ARE YOUR PRIMARY SOURCE FOR FINANCIALS.**
The user has uploaded analyst reports, models, and research files above. For Part 4 (all financial sections):
- Extract every number you can from the uploaded files FIRST before using SEC or Yahoo Finance data
- The uploaded Excel model is likely the most accurate and detailed financial model — use it as the authoritative source for revenue, margins, opex line items, EPS, capex, and DCF assumptions
- The uploaded PDFs likely contain analyst estimates, price targets, and thesis — extract specific figures, not summaries
- Where uploaded data conflicts with SEC/Yahoo Finance, prefer the uploaded data and note the discrepancy
- If the model has a full income statement breakdown (R&D, S&M, G&A), use those exact figures — do not leave line items as n/a if the uploaded model has them

Use the web_search tool liberally to:
- Fill any gaps in the data above
- Research competitors and industry dynamics
- Find recent earnings call highlights and management commentary
- Identify upcoming catalysts (next earnings date, conferences, regulatory)
- Check Glassdoor / LinkedIn for management / culture signals
- Find historical comparable setups in the market
- Get latest analyst price targets and thesis

Be specific — use numbers, percentages, and specific names. Avoid vague statements.
Where data is genuinely unavailable after searching, say so explicitly.

**CITE YOUR SOURCES THROUGHOUT.**
Every factual claim, number, target, or data point must have a source tag immediately after it. Use these formats:
- Uploaded file: [Source: Goldman Sachs initiation, Mar 2025] or [Source: uploaded model]
- SEC filing: [Source: 10-Q Q3 2024] or [Source: 8-K press release]
- Earnings call: [Source: Q4 2024 earnings call]
- Web search: [Source: Bloomberg, Mar 2025] or [Source: WSJ] or [Source: web search]
- yfinance data: [Source: yfinance]
- Stock Digest Telegram: [Source: Stock Digest]
If a figure appears in both an uploaded file and a web result, cite the uploaded file.
Do NOT write entire paragraphs without any source tags.

Start with Part 1 (Quick Pass/Fail). If this is an obvious fail, say so clearly and briefly.
Then complete all 10 parts of the framework.
"""

    messages = [{"role": "user", "content": prompt}]

    print(f"  Calling Claude Opus with extended thinking...")

    while True:
        response = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=16000,
            thinking={"type": "adaptive"},
            tools=[SEARCH_TOOL],
            messages=messages,
        )
        if response.stop_reason != "tool_use":
            break

        tool_blocks = [b for b in response.content if b.type == "tool_use"]
        print(f"  [web search x{len(tool_blocks)}] {', '.join(b.input['query'] for b in tool_blocks)}")

        def _search(block):
            return {
                "type": "tool_result",
                "tool_use_id": block.id,
                "content": web_search(block.input["query"]),
            }

        from concurrent.futures import ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=len(tool_blocks)) as pool:
            tool_results = list(pool.map(_search, tool_blocks))
        messages.append({"role": "assistant", "content": response.content})
        messages.append({"role": "user", "content": tool_results})

    report_text = ""
    for block in response.content:
        if block.type == "text":
            report_text = block.text.strip()

    return report_text


def export_docx(ticker: str, company_name: str, report_text: str, output_dir: str) -> str:
    from docx import Document
    from docx.shared import Pt
    import re

    date_str = datetime.now().strftime("%Y-%m-%d")
    safe_name = re.sub(r'[\\/:*?"<>|]', "", company_name).strip()
    filename = f"{ticker}_{safe_name}_{date_str}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_runs(para, text):
        """Add runs to a paragraph, bolding **text** spans."""
        for part in re.split(r"(\*\*[^*]+\*\*|__[^_]+__)", text):
            if (part.startswith("**") and part.endswith("**")) or \
               (part.startswith("__") and part.endswith("__")):
                para.add_run(part[2:-2]).bold = True
            else:
                para.add_run(part)

    def add_table(table_lines):
        """Parse markdown table lines and render as a Word table."""
        parsed = []
        for tl in table_lines:
            if re.match(r"^\s*\|[-| :]+\|\s*$", tl):
                continue  # skip separator rows
            cells = [c.strip() for c in tl.strip("|").split("|")]
            parsed.append(cells)
        if not parsed:
            return
        ncols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=ncols)
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(parsed):
            for c_idx in range(ncols):
                text = row[c_idx].replace("**", "").replace("__", "") if c_idx < len(row) else ""
                cell = tbl.cell(r_idx, c_idx)
                cell.text = ""
                run = cell.paragraphs[0].add_run(text)
                if r_idx == 0:
                    run.bold = True
        doc.add_paragraph()

    lines = report_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Markdown table block
        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(block)
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif not stripped or stripped in ("---", "==="):
            doc.add_paragraph()
        elif re.match(r"^[-*] ", stripped):
            add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\. ", "", stripped))
        else:
            add_runs(doc.add_paragraph(), stripped)

        i += 1

    doc.save(filepath)
    print(f"  Docx saved: {filepath}")
    return filepath


def export_pdf(docx_path: str) -> str:
    import platform, subprocess
    pdf_path = docx_path.replace(".docx", ".pdf")
    if platform.system() == "Windows":
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    else:
        output_dir = str(Path(docx_path).parent)
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            check=True, capture_output=True,
        )
    print(f"  PDF saved: {pdf_path}")
    return pdf_path
