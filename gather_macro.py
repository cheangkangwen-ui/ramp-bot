"""
Data gathering for macro ramp.
Pulls: yfinance price + technicals, asset-type-specific web search,
and user-uploaded PDFs/Excels. No SEC EDGAR or earnings transcripts.
"""

import os
import re
import pandas as pd
import yfinance as yf
from pathlib import Path

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

MAX_CHARS_PRICE    = 6000
MAX_CHARS_NEWS     = 10000
MAX_CHARS_FILE     = 8000
MAX_CHARS_TELEGRAM = 12000

TELEGRAM_API_ID   = 33919151
TELEGRAM_API_HASH = "dd0a935bd6545cf56910292ff4445c4e"

# ── Asset registry ──────────────────────────────────────────────────────────────

ASSET_MAP = {
    # ── Precious Metals ──────────────────────────────────────────────────────────
    "GOLD": "GC=F",       "XAU": "GC=F",
    "SILVER": "SI=F",     "XAG": "SI=F",
    "PLATINUM": "PL=F",   "PLAT": "PL=F",   "XPT": "PL=F",
    "PALLADIUM": "PA=F",  "PALL": "PA=F",   "XPD": "PA=F",
    "GLD": "GLD",         "SLV": "SLV",
    "GDX": "GDX",         "GDXJ": "GDXJ",

    # ── Energy ───────────────────────────────────────────────────────────────────
    "WTI": "CL=F",        "OIL": "CL=F",    "CRUDE": "CL=F",
    "BRENT": "BZ=F",      "BRENTOIL": "BZ=F",
    "NATGAS": "NG=F",     "GAS": "NG=F",    "NG": "NG=F",
    "RBOB": "RB=F",       "GASOLINE": "RB=F",
    "HEATINGOIL": "HO=F", "HO": "HO=F",
    "USO": "USO",         "UNG": "UNG",

    # ── Base Metals ──────────────────────────────────────────────────────────────
    "COPPER": "HG=F",     "HG": "HG=F",
    "ALUMINUM": "ALI=F",  "ALUMINIUM": "ALI=F",
    "ZINC": "^ZINC",
    "NICKEL": "^NICKEL",

    # ── Agricultural ─────────────────────────────────────────────────────────────
    "WHEAT": "ZW=F",
    "CORN": "ZC=F",
    "SOYBEAN": "ZS=F",    "SOY": "ZS=F",
    "SOYMEAL": "ZM=F",
    "SOYOIL": "ZL=F",
    "COFFEE": "KC=F",
    "SUGAR": "SB=F",
    "COCOA": "CC=F",
    "COTTON": "CT=F",
    "LIVECATTLE": "LE=F", "CATTLE": "LE=F",
    "LEANHOGS": "HE=F",   "HOGS": "HE=F",
    "LUMBER": "LBR=F",
    "OJ": "OJ=F",         "ORANGEJUICE": "OJ=F",
    "DBA": "DBA",

    # ── US Rates ─────────────────────────────────────────────────────────────────
    "US2Y": "^IRX",       "2Y": "^IRX",     "TBILL": "^IRX",
    "US5Y": "^FVX",       "5Y": "^FVX",
    "US10Y": "^TNX",      "10Y": "^TNX",    "TNX": "^TNX",
    "US30Y": "^TYX",      "30Y": "^TYX",    "TYX": "^TYX",
    "TLT": "TLT",         "IEF": "IEF",     "SHY": "SHY",
    "TBT": "TBT",         "TMF": "TMF",     "TMV": "TMV",
    "HYG": "HYG",         "JNK": "JNK",
    "LQD": "LQD",         "EMB": "EMB",
    "TIPS": "TIP",        "TIP": "TIP",

    # ── Global Rates / Bonds ─────────────────────────────────────────────────────
    "GILT": "^TNX",       "UK10Y": "^TNX",   # closest proxy — use directly for now
    "BUND": "^DE10YB-EUR", "GERMAN10Y": "^DE10YB-EUR",
    "JGB": "^JP10YB-JPN",  "JAPAN10Y": "^JP10YB-JPN",

    # ── G10 FX ───────────────────────────────────────────────────────────────────
    "EURUSD": "EURUSD=X",  "EUR": "EURUSD=X",
    "USDJPY": "JPY=X",     "JPY": "JPY=X",
    "GBPUSD": "GBPUSD=X",  "GBP": "GBPUSD=X",  "CABLE": "GBPUSD=X",
    "AUDUSD": "AUDUSD=X",  "AUD": "AUDUSD=X",  "AUSSIE": "AUDUSD=X",
    "NZDUSD": "NZDUSD=X",  "NZD": "NZDUSD=X",  "KIWI": "NZDUSD=X",
    "USDCAD": "CAD=X",     "CAD": "CAD=X",     "LOONIE": "CAD=X",
    "USDCHF": "CHF=X",     "CHF": "CHF=X",     "SWISSY": "CHF=X",
    "DXY": "DX-Y.NYB",     "DOLLAR": "DX-Y.NYB", "USD": "DX-Y.NYB",
    "UUP": "UUP",

    # ── Cross Rates ──────────────────────────────────────────────────────────────
    "EURJPY": "EURJPY=X",
    "EURGBP": "EURGBP=X",
    "GBPJPY": "GBPJPY=X",
    "AUDJPY": "AUDJPY=X",
    "EURAUD": "EURAUD=X",
    "EURCHF": "EURCHF=X",
    "GBPAUD": "GBPAUD=X",

    # ── EM FX ────────────────────────────────────────────────────────────────────
    "USDCNY": "CNY=X",    "CNY": "CNY=X",   "YUAN": "CNY=X",
    "USDCNH": "CNH=X",    "CNH": "CNH=X",
    "USDINR": "INR=X",    "INR": "INR=X",   "RUPEE": "INR=X",
    "USDBRL": "BRL=X",    "BRL": "BRL=X",   "REAL": "BRL=X",
    "USDMXN": "MXN=X",    "MXN": "MXN=X",   "PESO": "MXN=X",
    "USDZAR": "ZAR=X",    "ZAR": "ZAR=X",   "RAND": "ZAR=X",
    "USDTRY": "TRY=X",    "TRY": "TRY=X",   "LIRA": "TRY=X",
    "USDKRW": "KRW=X",    "KRW": "KRW=X",   "WON": "KRW=X",
    "USDSGD": "SGD=X",    "SGD": "SGD=X",
    "USDHKD": "HKD=X",    "HKD": "HKD=X",
    "USDTHB": "THB=X",    "THB": "THB=X",
    "USDIDR": "IDR=X",    "IDR": "IDR=X",
    "USDPLN": "PLN=X",    "PLN": "PLN=X",
    "USDSEK": "SEK=X",    "SEK": "SEK=X",
    "USDNOK": "NOK=X",    "NOK": "NOK=X",

    # ── US Equity Indexes ────────────────────────────────────────────────────────
    "SPX": "^GSPC",       "SP500": "^GSPC",  "ES": "^GSPC",
    "NDX": "^NDX",        "NASDAQ": "^NDX",  "NQ": "^NDX",
    "DOW": "^DJI",        "DJIA": "^DJI",
    "RTY": "^RUT",        "RUSSELL": "^RUT", "RUSSELL2000": "^RUT",
    "VIX": "^VIX",
    "SPY": "SPY",         "QQQ": "QQQ",      "IWM": "IWM",
    "DIA": "DIA",

    # ── European Indexes ─────────────────────────────────────────────────────────
    "FTSE": "^FTSE",      "FTSE100": "^FTSE", "UKX": "^FTSE",
    "DAX": "^GDAXI",      "DAX40": "^GDAXI",
    "CAC": "^FCHI",       "CAC40": "^FCHI",
    "IBEX": "^IBEX",      "IBEX35": "^IBEX",
    "SMI": "^SSMI",
    "AEX": "^AEX",
    "MIB": "FTSEMIB.MI",  "FTMIB": "FTSEMIB.MI",
    "STOXX50": "^STOXX50E", "EUROSTOXX": "^STOXX50E",
    "VSTOXX": "^V2TX",    "VSTOX": "^V2TX",

    # ── Asia-Pacific Indexes ─────────────────────────────────────────────────────
    "NIKKEI": "^N225",    "NKY": "^N225",    "N225": "^N225",
    "TOPIX": "^TOPX",
    "HSI": "^HSI",        "HANGSENG": "^HSI",
    "CSI300": "000300.SS", "A50": "000300.SS",
    "KOSPI": "^KS11",     "KS11": "^KS11",
    "ASX200": "^AXJO",    "ASX": "^AXJO",
    "SENSEX": "^BSESN",   "NIFTY": "^NSEI",  "NIFTY50": "^NSEI",
    "TWSE": "^TWII",      "TAIWAN": "^TWII",
    "STI": "^STI",        "SINGAPORE": "^STI",

    # ── EM / Other Indexes ───────────────────────────────────────────────────────
    "BOVESPA": "^BVSP",   "IBOV": "^BVSP",   "BRAZIL": "^BVSP",
    "MEXBOL": "^MXX",     "MEXICO": "^MXX",
    "EEM": "EEM",         "VWO": "VWO",       "EWZ": "EWZ",

    # ── Crypto ───────────────────────────────────────────────────────────────────
    "BTC": "BTC-USD",     "BITCOIN": "BTC-USD",
    "ETH": "ETH-USD",     "ETHEREUM": "ETH-USD",
    "SOL": "SOL-USD",     "SOLANA": "SOL-USD",
    "XRP": "XRP-USD",
    "IBIT": "IBIT",       "FBTC": "FBTC",
}

ASSET_TYPES = {
    "commodity": {
        "GC=F", "SI=F", "PL=F", "PA=F", "GLD", "SLV", "GDX", "GDXJ",
        "CL=F", "BZ=F", "NG=F", "RB=F", "HO=F", "USO", "UNG",
        "HG=F", "ALI=F", "^ZINC", "^NICKEL",
        "ZW=F", "ZC=F", "ZS=F", "ZM=F", "ZL=F", "KC=F", "SB=F",
        "CC=F", "CT=F", "LE=F", "HE=F", "LBR=F", "OJ=F", "DBA",
    },
    "rates": {
        "^IRX", "^FVX", "^TNX", "^TYX",
        "TLT", "IEF", "SHY", "TBT", "TMF", "TMV",
        "HYG", "JNK", "LQD", "EMB", "TIP",
        "^DE10YB-EUR", "^JP10YB-JPN",
    },
    "fx": {
        "EURUSD=X", "JPY=X", "GBPUSD=X", "AUDUSD=X", "NZDUSD=X",
        "CAD=X", "CHF=X", "DX-Y.NYB", "UUP",
        "EURJPY=X", "EURGBP=X", "GBPJPY=X", "AUDJPY=X", "EURAUD=X",
        "EURCHF=X", "GBPAUD=X",
        "CNY=X", "CNH=X", "INR=X", "BRL=X", "MXN=X", "ZAR=X",
        "TRY=X", "KRW=X", "SGD=X", "HKD=X", "THB=X", "IDR=X",
        "PLN=X", "SEK=X", "NOK=X",
    },
    "equity": {
        "^GSPC", "^NDX", "^DJI", "^RUT", "^VIX",
        "SPY", "QQQ", "IWM", "DIA",
        "^FTSE", "^GDAXI", "^FCHI", "^IBEX", "^SSMI", "^AEX",
        "FTSEMIB.MI", "^STOXX50E", "^V2TX",
        "^N225", "^TOPX", "^HSI", "000300.SS", "^KS11", "^AXJO",
        "^BSESN", "^NSEI", "^TWII", "^STI",
        "^BVSP", "^MXX", "EEM", "VWO", "EWZ",
    },
    "crypto": {
        "BTC-USD", "ETH-USD", "SOL-USD", "XRP-USD", "IBIT", "FBTC",
    },
}

ASSET_DISPLAY_NAMES = {
    # Precious metals
    "GC=F": "Gold", "SI=F": "Silver", "PL=F": "Platinum", "PA=F": "Palladium",
    "GLD": "SPDR Gold ETF", "SLV": "iShares Silver ETF",
    "GDX": "VanEck Gold Miners ETF", "GDXJ": "VanEck Junior Gold Miners ETF",
    # Energy
    "CL=F": "WTI Crude Oil", "BZ=F": "Brent Crude",
    "NG=F": "Natural Gas", "RB=F": "RBOB Gasoline", "HO=F": "Heating Oil",
    "USO": "United States Oil ETF", "UNG": "United States Natural Gas ETF",
    # Base metals
    "HG=F": "Copper", "ALI=F": "Aluminum", "^ZINC": "Zinc", "^NICKEL": "Nickel",
    # Agricultural
    "ZW=F": "Wheat", "ZC=F": "Corn", "ZS=F": "Soybean", "ZM=F": "Soybean Meal",
    "ZL=F": "Soybean Oil", "KC=F": "Coffee", "SB=F": "Sugar", "CC=F": "Cocoa",
    "CT=F": "Cotton", "LE=F": "Live Cattle", "HE=F": "Lean Hogs",
    "LBR=F": "Lumber", "OJ=F": "Orange Juice", "DBA": "Invesco Agri ETF",
    # US Rates
    "^IRX": "US 2Y Treasury Yield", "^FVX": "US 5Y Treasury Yield",
    "^TNX": "US 10Y Treasury Yield", "^TYX": "US 30Y Treasury Yield",
    "TLT": "iShares 20Y+ Treasury ETF", "IEF": "iShares 7-10Y Treasury ETF",
    "SHY": "iShares 1-3Y Treasury ETF", "TBT": "ProShares Short 20Y Treasury",
    "TMF": "Direxion 3x Long Treasury", "TMV": "Direxion 3x Short Treasury",
    "HYG": "iShares HY Bond ETF", "JNK": "SPDR HY Bond ETF",
    "LQD": "iShares IG Bond ETF", "EMB": "iShares EM Bond ETF", "TIP": "iShares TIPS ETF",
    # Global Rates
    "^DE10YB-EUR": "Germany 10Y Bund Yield", "^JP10YB-JPN": "Japan 10Y JGB Yield",
    # G10 FX
    "EURUSD=X": "EUR/USD", "JPY=X": "USD/JPY", "GBPUSD=X": "GBP/USD",
    "AUDUSD=X": "AUD/USD", "NZDUSD=X": "NZD/USD", "CAD=X": "USD/CAD",
    "CHF=X": "USD/CHF", "DX-Y.NYB": "DXY (Dollar Index)", "UUP": "Invesco Dollar ETF",
    # Cross rates
    "EURJPY=X": "EUR/JPY", "EURGBP=X": "EUR/GBP", "GBPJPY=X": "GBP/JPY",
    "AUDJPY=X": "AUD/JPY", "EURAUD=X": "EUR/AUD", "EURCHF=X": "EUR/CHF",
    "GBPAUD=X": "GBP/AUD",
    # EM FX
    "CNY=X": "USD/CNY", "CNH=X": "USD/CNH", "INR=X": "USD/INR",
    "BRL=X": "USD/BRL", "MXN=X": "USD/MXN", "ZAR=X": "USD/ZAR",
    "TRY=X": "USD/TRY", "KRW=X": "USD/KRW", "SGD=X": "USD/SGD",
    "HKD=X": "USD/HKD", "THB=X": "USD/THB", "IDR=X": "USD/IDR",
    "PLN=X": "USD/PLN", "SEK=X": "USD/SEK", "NOK=X": "USD/NOK",
    # US Equity
    "^GSPC": "S&P 500", "^NDX": "Nasdaq 100", "^DJI": "Dow Jones",
    "^RUT": "Russell 2000", "^VIX": "VIX",
    "SPY": "SPDR S&P 500 ETF", "QQQ": "Invesco QQQ ETF",
    "IWM": "iShares Russell 2000 ETF", "DIA": "SPDR Dow Jones ETF",
    # European Indexes
    "^FTSE": "FTSE 100", "^GDAXI": "DAX", "^FCHI": "CAC 40",
    "^IBEX": "IBEX 35", "^SSMI": "SMI", "^AEX": "AEX",
    "FTSEMIB.MI": "FTSE MIB", "^STOXX50E": "Euro Stoxx 50", "^V2TX": "VSTOXX",
    # Asia-Pacific Indexes
    "^N225": "Nikkei 225", "^TOPX": "TOPIX", "^HSI": "Hang Seng",
    "000300.SS": "CSI 300", "^KS11": "KOSPI", "^AXJO": "ASX 200",
    "^BSESN": "BSE Sensex", "^NSEI": "Nifty 50", "^TWII": "TWSE",
    "^STI": "Straits Times Index",
    # EM / Other Indexes
    "^BVSP": "Bovespa", "^MXX": "IPC Mexico",
    "EEM": "iShares MSCI EM ETF", "VWO": "Vanguard FTSE EM ETF",
    "EWZ": "iShares MSCI Brazil ETF",
    # Crypto
    "BTC-USD": "Bitcoin", "ETH-USD": "Ethereum", "SOL-USD": "Solana",
    "XRP-USD": "XRP", "IBIT": "iShares Bitcoin ETF", "FBTC": "Fidelity Bitcoin ETF",
}


def resolve_asset(name: str) -> tuple[str, str, str]:
    """Return (yf_ticker, display_name, asset_type) for a given input name."""
    key = name.upper().replace("/", "").replace("-", "").replace(".", "")
    yf_ticker = ASSET_MAP.get(key, name)
    display_name = ASSET_DISPLAY_NAMES.get(yf_ticker, yf_ticker)
    asset_type = "unknown"
    for atype, tickers in ASSET_TYPES.items():
        if yf_ticker in tickers:
            asset_type = atype
            break
    return yf_ticker, display_name, asset_type


# ── Web search ──────────────────────────────────────────────────────────────────

def web_search(query: str, max_results: int = 6) -> str:
    try:
        from ddgs import DDGS
    except ImportError:
        from duckduckgo_search import DDGS
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
        if not results:
            return "No results found."
        return "\n\n".join(
            f"**{r['title']}**\n{r['href']}\n{r['body']}" for r in results
        )
    except Exception as e:
        return f"Search failed: {e}"


def gather_news(display_name: str, asset_type: str) -> str:
    if asset_type == "commodity":
        queries = [
            f"{display_name} price outlook supply demand 2025",
            f"{display_name} inventory stockpiles COT CFTC positioning",
            f"{display_name} macro drivers central bank geopolitical",
            f"{display_name} technical analysis support resistance analyst target",
        ]
    elif asset_type == "rates":
        queries = [
            f"{display_name} yield outlook Federal Reserve ECB policy 2025",
            f"{display_name} inflation expectations real yield breakeven",
            f"{display_name} Treasury supply foreign demand COT positioning",
            f"{display_name} technical levels support resistance analyst forecast",
        ]
    elif asset_type == "fx":
        queries = [
            f"{display_name} rate differential central bank outlook 2025",
            f"{display_name} COT IMM speculative positioning net long short",
            f"{display_name} technical analysis support resistance levels",
            f"{display_name} PPP fair value current account inflation differential",
        ]
    elif asset_type == "equity":
        queries = [
            f"{display_name} index outlook earnings valuation 2025",
            f"{display_name} COT positioning put call ratio sentiment",
            f"{display_name} technical analysis support resistance",
            f"{display_name} macro risks Fed liquidity earnings recession",
        ]
    elif asset_type == "crypto":
        queries = [
            f"{display_name} price outlook on-chain data 2025",
            f"{display_name} institutional flows ETF inflows positioning",
            f"{display_name} technical analysis support resistance levels",
            f"{display_name} macro correlation risk-on risk-off Fed liquidity",
        ]
    else:
        queries = [
            f"{display_name} price outlook 2025",
            f"{display_name} positioning sentiment COT",
            f"{display_name} technical analysis levels",
            f"{display_name} macro drivers risks",
        ]

    parts = []
    for q in queries:
        result = web_search(q, max_results=5)
        parts.append(f"### Search: {q}\n{result}")
    full = "\n\n".join(parts)
    if len(full) > MAX_CHARS_NEWS:
        full = full[:MAX_CHARS_NEWS] + "\n\n[... truncated ...]"
    return full


# ── Price data + technicals ─────────────────────────────────────────────────────

def gather_price_data(yf_ticker: str, display_name: str) -> str:
    t = yf.Ticker(yf_ticker)
    parts = [f"## {display_name} ({yf_ticker}) — Price & Technicals"]

    try:
        hist = t.history(period="2y")
        if hist.empty:
            return f"No price data available for {yf_ticker}."

        close = hist["Close"].dropna()
        high_series = hist["High"].dropna()
        low_series = hist["Low"].dropna()
        current = close.iloc[-1]

        # Moving averages
        ma14  = close.rolling(14).mean().iloc[-1]
        ma50  = close.rolling(50).mean().iloc[-1]
        ma200 = close.rolling(200).mean().iloc[-1]

        # RSI(14)
        delta = close.diff()
        gain = delta.clip(lower=0).rolling(14).mean()
        loss = (-delta).clip(lower=0).rolling(14).mean()
        rsi = (100 - 100 / (1 + gain / loss)).iloc[-1]

        # ATR(20)
        prev_close = close.shift(1)
        tr = pd.concat([
            high_series - low_series,
            (high_series - prev_close).abs(),
            (low_series - prev_close).abs(),
        ], axis=1).max(axis=1)
        atr20 = tr.rolling(20).mean().iloc[-1]

        # 52-week range
        high52 = close.rolling(252).max().iloc[-1]
        low52  = close.rolling(252).min().iloc[-1]
        pct_from_high = (current / high52 - 1) * 100
        pct_from_low  = (current / low52  - 1) * 100

        # YTD
        ytd_start = close[close.index.year == close.index[-1].year].iloc[0]
        ytd_change = (current / ytd_start - 1) * 100

        # 1M, 3M, 6M, 1Y returns
        def ret(n_days):
            if len(close) > n_days:
                return (current / close.iloc[-n_days] - 1) * 100
            return None

        parts.append(f"\n### Current Price: {current:.4f}")
        parts.append(
            f"52W High: {high52:.4f} ({pct_from_high:+.1f}% from high)\n"
            f"52W Low:  {low52:.4f} ({pct_from_low:+.1f}% from low)\n"
            f"YTD: {ytd_change:+.1f}%\n"
            f"1M: {ret(21):+.1f}%  |  3M: {ret(63):+.1f}%  |  6M: {ret(126):+.1f}%  |  1Y: {ret(252):+.1f}%"
        )

        parts.append(
            f"\n### Moving Averages\n"
            f"14d MA:  {ma14:.4f}  ({'ABOVE' if current > ma14 else 'BELOW'})\n"
            f"50d MA:  {ma50:.4f}  ({'ABOVE' if current > ma50 else 'BELOW'})\n"
            f"200d MA: {ma200:.4f} ({'ABOVE' if current > ma200 else 'BELOW'})"
        )

        ma_alignment = "Bullish (14>50>200)" if ma14 > ma50 > ma200 else \
                       "Bearish (14<50<200)" if ma14 < ma50 < ma200 else "Mixed"
        parts.append(f"MA Alignment: {ma_alignment}")

        parts.append(
            f"\n### Momentum\n"
            f"RSI(14): {rsi:.1f}  "
            f"({'Overbought >70' if rsi > 70 else 'Oversold <30' if rsi < 30 else 'Neutral'})\n"
            f"ATR(20): {atr20:.4f} ({(atr20/current)*100:.2f}% of price)"
        )

        # Recent weekly closes
        weekly = close.resample("W").last().tail(26)
        parts.append(f"\n### Weekly Closes (last 26 weeks)\n{weekly.to_string()}")

    except Exception as e:
        parts.append(f"Price data error: {e}")

    result = "\n".join(parts)
    if len(result) > MAX_CHARS_PRICE:
        result = result[:MAX_CHARS_PRICE] + "\n\n[... truncated ...]"
    return result


# ── File parsing ────────────────────────────────────────────────────────────────

def parse_pdf(file_path: str) -> str:
    if not HAS_PDFPLUMBER:
        return "[pdfplumber not installed — PDF parsing unavailable]"
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        full_text = "\n\n".join(text_parts).strip()
        if len(full_text) > MAX_CHARS_FILE:
            full_text = full_text[:MAX_CHARS_FILE] + "\n\n[... truncated ...]"
        return full_text
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def parse_excel(file_path: str) -> str:
    try:
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path)
            result = df.to_string(max_rows=100)
        else:
            xl = pd.ExcelFile(file_path)
            parts = []
            for sheet in xl.sheet_names[:5]:
                df = xl.parse(sheet)
                parts.append(f"### Sheet: {sheet}\n{df.to_string(max_rows=80, max_cols=20)}")
            result = "\n\n".join(parts)
        if len(result) > MAX_CHARS_FILE:
            result = result[:MAX_CHARS_FILE] + "\n\n[... truncated ...]"
        return result
    except Exception as e:
        return f"[Excel/CSV extraction failed: {e}]"


def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                text_parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        full_text = "\n\n".join(text_parts).strip()
        if len(full_text) > MAX_CHARS_FILE:
            full_text = full_text[:MAX_CHARS_FILE] + "\n\n[... truncated ...]"
        return full_text
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def parse_uploaded_files(file_paths: list, captions: dict = None) -> str:
    if not file_paths:
        return ""
    captions = captions or {}
    parts = []
    for fp in file_paths:
        if not os.path.exists(fp):
            continue
        ext = Path(fp).suffix.lower()
        name = Path(fp).name
        if ext == ".pdf":
            content = parse_pdf(fp)
        elif ext in (".docx", ".doc"):
            content = parse_docx(fp)
        elif ext in (".xlsx", ".xls", ".csv"):
            content = parse_excel(fp)
        else:
            content = f"[Unsupported file type: {ext}]"
        header = f"## Uploaded File: {name}"
        if fp in captions and captions[fp]:
            header += f"\nUser context: {captions[fp]}"
        parts.append(f"{header}\n{content}")
    return "\n\n".join(parts)


# ── Telegram digest ─────────────────────────────────────────────────────────────

async def _fetch_telegram_messages(group_name: str, search_query: str, days: int, max_messages: int) -> str:
    from telethon import TelegramClient
    from telethon.sessions import StringSession
    from datetime import datetime, timedelta, timezone
    session_str = os.environ.get("TELEGRAM_SESSION", "")
    session = StringSession(session_str) if len(session_str) > 20 else session_str
    tg = TelegramClient(session, TELEGRAM_API_ID, TELEGRAM_API_HASH)
    try:
        await tg.connect()
        if not await tg.is_user_authorized():
            return "Telegram session not authorized."
        dialogs = await tg.get_dialogs()
        group_entity = None
        for d in dialogs:
            if d.name == group_name:
                group_entity = d.entity
                break
        if group_entity is None:
            return f"Group '{group_name}' not found."
        cutoff = datetime.now(timezone.utc) - timedelta(days=days)
        messages = []
        async for msg in tg.iter_messages(group_entity, search=search_query, limit=max_messages):
            if msg.date < cutoff:
                break
            if msg.text and msg.text.strip():
                messages.append(f"[{msg.date.strftime('%Y-%m-%d')}] {msg.text.strip()}")
        if not messages:
            return f"No messages found for '{search_query}' in '{group_name}' (last {days} days)."
        result = f"## {group_name} — '{search_query}' ({len(messages)} messages, last {days} days)\n\n"
        result += "\n\n---\n\n".join(messages)
        if len(result) > MAX_CHARS_TELEGRAM:
            result = result[:MAX_CHARS_TELEGRAM] + "\n\n[... truncated ...]"
        return result
    finally:
        await tg.disconnect()


def gather_telegram_digest(group_name: str, search_query: str, days: int, max_messages: int = 50) -> str:
    import asyncio
    try:
        return asyncio.run(_fetch_telegram_messages(group_name, search_query, days, max_messages))
    except Exception as e:
        return f"Telegram digest fetch failed: {e}"


# ── Main entry point ────────────────────────────────────────────────────────────

def gather_all(asset_input: str, staged_files: list, staged_captions: dict = None) -> dict:
    """
    Gather all data for a macro asset. Returns dict with keys:
    yf_ticker, display_name, asset_type, price_data, news, uploaded_files
    """
    yf_ticker, display_name, asset_type = resolve_asset(asset_input)

    print(f"  Asset: {display_name} ({yf_ticker}) | Type: {asset_type}")

    print(f"  Gathering price data + technicals...")
    price_data = gather_price_data(yf_ticker, display_name)

    print(f"  Searching macro news and context...")
    news_data = gather_news(display_name, asset_type)

    print(f"  Parsing {len(staged_files)} staged file(s)...")
    uploaded = parse_uploaded_files(staged_files, staged_captions)

    print(f"  Searching News Digest Telegram group...")
    telegram_data = gather_telegram_digest("📊 News Digest", display_name, days=30, max_messages=50)

    return {
        "asset_input":    asset_input,
        "yf_ticker":      yf_ticker,
        "display_name":   display_name,
        "asset_type":     asset_type,
        "price_data":     price_data,
        "news":           news_data,
        "uploaded_files": uploaded,
        "telegram_digest": telegram_data,
    }
