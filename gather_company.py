"""
Data gathering for company ramp.
Pulls: yfinance financials/stats, SEC EDGAR filings (actual content),
DuckDuckGo news search, earnings transcript (Quartr → SEC exhibit → web),
and user-uploaded PDFs/Excels.
"""

import os
import re
import json
import requests
import pdfplumber
import yfinance as yf
import pandas as pd
from pathlib import Path

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

MAX_CHARS_YFINANCE = 10000
MAX_CHARS_NEWS = 8000
MAX_CHARS_PRESS_RELEASE = 8000
MAX_CHARS_10Q = 8000
MAX_CHARS_10K = 5000
MAX_CHARS_TRANSCRIPT = 14000
MAX_CHARS_FILE = 8000

EDGAR_HEADERS = {"User-Agent": "company-ramp-bot kangw@personal.com"}
QUARTR_API_BASE = "https://api.quartr.com/public/v3"


# ── Web search ─────────────────────────────────────────────────────────────────

def web_search(query: str, max_results: int = 6) -> str:
    try:
        from ddgs import DDGS
    except ImportError:
        from duckduckgo_search import DDGS
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
        if not results:
            return "No results found."
        return "\n\n".join(
            f"**{r['title']}**\n{r['href']}\n{r['body']}" for r in results
        )
    except Exception as e:
        return f"Search failed: {e}"


def gather_news(ticker: str, company_name: str) -> str:
    queries = [
        f"{ticker} {company_name} stock news earnings 2025",
        f"{ticker} {company_name} analyst rating price target",
        f"{company_name} industry competition market share outlook",
        f"{ticker} management CEO strategy recent",
    ]
    parts = []
    for q in queries:
        result = web_search(q, max_results=5)
        parts.append(f"### Search: {q}\n{result}")
    full = "\n\n".join(parts)
    if len(full) > MAX_CHARS_NEWS:
        full = full[:MAX_CHARS_NEWS] + "\n\n[... truncated ...]"
    return full


# ── Yahoo Finance ──────────────────────────────────────────────────────────────

def gather_yfinance(ticker: str) -> str:
    t = yf.Ticker(ticker)
    parts = []

    try:
        info = t.info
        parts.append("## Company Info")
        for key in [
            "longName", "sector", "industry", "longBusinessSummary",
            "fullTimeEmployees", "country", "website",
            "marketCap", "enterpriseValue",
        ]:
            val = info.get(key)
            if val:
                parts.append(f"{key}: {val}")

        parts.append("\n## Key Statistics")
        for key in [
            "currentPrice", "fiftyTwoWeekHigh", "fiftyTwoWeekLow",
            "fiftyDayAverage", "twoHundredDayAverage",
            "trailingPE", "forwardPE", "priceToBook", "priceToSalesTrailing12Months",
            "enterpriseToEbitda", "enterpriseToRevenue",
            "beta", "52WeekChange",
            "shortPercentOfFloat", "shortRatio",
            "heldPercentInstitutions", "heldPercentInsiders",
            "dividendYield", "payoutRatio",
            "returnOnEquity", "returnOnAssets",
            "profitMargins", "operatingMargins", "grossMargins",
            "revenueGrowth", "earningsGrowth",
            "totalDebt", "totalCash", "debtToEquity", "currentRatio",
            "targetHighPrice", "targetLowPrice", "targetMeanPrice",
            "numberOfAnalystOpinions", "recommendationMean", "recommendationKey",
        ]:
            val = info.get(key)
            if val is not None:
                parts.append(f"{key}: {val}")
    except Exception as e:
        parts.append(f"Info fetch failed: {e}")

    for label, attr in [
        ("Income Statement (Annual)", "income_stmt"),
        ("Balance Sheet (Annual)", "balance_sheet"),
        ("Cash Flow (Annual)", "cash_flow"),
    ]:
        try:
            df = getattr(t, attr)
            if df is not None and not df.empty:
                parts.append(f"\n## {label}")
                parts.append(df.to_string())
        except Exception:
            pass

    try:
        eh = t.earnings_history
        if eh is not None and not eh.empty:
            parts.append("\n## Earnings History (last 8 quarters)")
            parts.append(eh.tail(8).to_string())
    except Exception:
        pass

    try:
        recs = t.recommendations
        if recs is not None and not recs.empty:
            parts.append("\n## Analyst Recommendations (recent)")
            parts.append(recs.tail(10).to_string())
    except Exception:
        pass

    try:
        cal = t.calendar
        if cal:
            parts.append("\n## Earnings Calendar")
            parts.append(str(cal))
    except Exception:
        pass

    try:
        hist = t.history(period="6mo")
        if not hist.empty:
            parts.append("\n## Price History (6mo, weekly close)")
            weekly = hist["Close"].resample("W").last()
            parts.append(weekly.to_string())
    except Exception:
        pass

    result = "\n".join(parts)
    if len(result) > MAX_CHARS_YFINANCE:
        result = result[:MAX_CHARS_YFINANCE] + "\n\n[... truncated ...]"
    return result


# ── SEC EDGAR helpers ──────────────────────────────────────────────────────────

def get_cik(ticker: str) -> str | None:
    """Look up CIK from EDGAR company ticker mapping."""
    try:
        resp = requests.get(
            "https://www.sec.gov/files/company_tickers.json",
            headers=EDGAR_HEADERS, timeout=15,
        )
        if resp.status_code == 200:
            for entry in resp.json().values():
                if entry.get("ticker", "").upper() == ticker.upper():
                    return str(entry["cik_str"]).zfill(10)
    except Exception:
        pass
    return None


def _fetch_limited(url: str, max_bytes: int = 400_000) -> bytes | None:
    """Download a URL, stopping after max_bytes to avoid huge filings."""
    try:
        resp = requests.get(url, headers=EDGAR_HEADERS, timeout=30, stream=True)
        if resp.status_code != 200:
            return None
        chunks = []
        total = 0
        for chunk in resp.iter_content(chunk_size=16384):
            chunks.append(chunk)
            total += len(chunk)
            if total >= max_bytes:
                break
        return b"".join(chunks)
    except Exception:
        return None


def _html_to_text(html: str) -> str:
    """Convert HTML to readable plain text, preserving table structure."""
    if HAS_BS4:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "head"]):
            tag.decompose()
        # Render tables as pipe-delimited rows
        for table in soup.find_all("table"):
            rows_text = []
            for tr in table.find_all("tr"):
                cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
                if any(cells):
                    rows_text.append(" | ".join(cells))
            table.replace_with("\n".join(rows_text) + "\n")
        text = soup.get_text(separator="\n", strip=True)
    else:
        text = re.sub(r"<[^>]+>", " ", html)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&amp;", "&", text)
        text = re.sub(r"&lt;", "<", text)
        text = re.sub(r"&gt;", ">", text)
        text = re.sub(r"&#\d+;", " ", text)

    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _find_financial_section(text: str) -> str:
    """Seek to the income statement / results section within a stripped filing."""
    markers = [
        "CONDENSED CONSOLIDATED STATEMENTS OF OPERATIONS",
        "CONSOLIDATED STATEMENTS OF OPERATIONS",
        "CONSOLIDATED STATEMENTS OF INCOME",
        "RESULTS OF OPERATIONS",
        "FINANCIAL HIGHLIGHTS",
        "SELECTED FINANCIAL DATA",
    ]
    upper = text.upper()
    best = len(text)
    for m in markers:
        pos = upper.find(m)
        if 0 <= pos < best:
            best = pos
    return text[best:] if best < len(text) else text


def _get_filing_docs(cik: str, accession: str) -> list[tuple[str, str]]:
    """Return list of (filename, url) from the filing's EDGAR index page."""
    acc_clean = accession.replace("-", "")
    base = f"https://www.sec.gov/Archives/edgar/data/{int(cik)}/{acc_clean}/"
    try:
        resp = requests.get(base, headers=EDGAR_HEADERS, timeout=15)
        if resp.status_code != 200:
            return []
        docs = []
        for m in re.finditer(r'href="([^"?#][^"]*)"', resp.text):
            href = m.group(1)
            if not href.startswith("/") and not href.startswith("http"):
                docs.append((href, base + href))
        return docs
    except Exception:
        return []


def _find_exhibit(docs: list, patterns: list) -> tuple[str | None, str | None]:
    """Find first document whose filename matches any regex pattern."""
    for filename, url in docs:
        for pat in patterns:
            if re.search(pat, filename, re.IGNORECASE):
                return filename, url
    return None, None


# ── SEC EDGAR — download actual filing content ─────────────────────────────────

def gather_sec(ticker: str) -> str:
    cik = get_cik(ticker)
    if not cik:
        return f"Could not find SEC CIK for {ticker}. Ticker may be non-US or delisted."

    try:
        resp = requests.get(
            f"https://data.sec.gov/submissions/CIK{cik}.json",
            headers=EDGAR_HEADERS, timeout=15,
        )
        meta = resp.json()
    except Exception as e:
        return f"SEC EDGAR fetch failed: {e}"

    company_name = meta.get("name", ticker)
    recent = meta.get("filings", {}).get("recent", {})
    forms = recent.get("form", [])
    accessions = recent.get("accessionNumber", [])
    dates = recent.get("filingDate", [])
    primary_docs = recent.get("primaryDocument", [])

    parts = [f"# SEC EDGAR — {company_name} (CIK: {cik})"]

    targets = [
        ("8-K",  MAX_CHARS_PRESS_RELEASE, "press_release"),
        ("10-Q", MAX_CHARS_10Q,           "quarterly"),
        ("10-K", MAX_CHARS_10K,           "annual"),
    ]

    for target_form, max_chars, mode in targets:
        for i, form in enumerate(forms[:40]):
            if form != target_form:
                continue

            acc = accessions[i]
            date = dates[i]
            primary_doc = primary_docs[i]
            acc_clean = acc.replace("-", "")
            cik_int = int(cik)

            parts.append(f"\n## {target_form} — filed {date}")

            try:
                if mode == "press_release":
                    # Discover the press release exhibit (ex99.1) within the 8-K
                    docs = _get_filing_docs(cik, acc)
                    pr_name, pr_url = _find_exhibit(docs, [
                        r"ex99.*earnings", r"ex99.*press", r"ex99.*result",
                        r"ex991\b", r"ex-991\b", r"exhibit.?99",
                        r"pressrelease", r"press.?release",
                    ])
                    if pr_url:
                        fetch_url = pr_url
                        parts.append(f"Source: press release exhibit ({pr_name})")
                    else:
                        fetch_url = f"https://www.sec.gov/Archives/edgar/data/{cik_int}/{acc_clean}/{primary_doc}"
                        parts.append(f"Source: primary 8-K document ({primary_doc}) — no exhibit found")
                else:
                    fetch_url = f"https://www.sec.gov/Archives/edgar/data/{cik_int}/{acc_clean}/{primary_doc}"
                    parts.append(f"Source: {primary_doc}")

                raw = _fetch_limited(fetch_url, max_bytes=500_000)
                if raw is None:
                    parts.append("[Download failed]")
                    break

                decoded = raw.decode("utf-8", errors="ignore")
                if "<html" in decoded.lower() or primary_doc.lower().endswith(".htm"):
                    text = _html_to_text(decoded)
                else:
                    text = decoded

                if mode in ("quarterly", "annual"):
                    text = _find_financial_section(text)

                text = re.sub(r"\n{3,}", "\n\n", text).strip()
                if len(text) > max_chars:
                    text = text[:max_chars] + "\n\n[... truncated ...]"

                parts.append(text)

            except Exception as e:
                parts.append(f"[Error: {e}]")

            break  # most recent filing only

    return "\n\n".join(parts)


# ── Earnings transcript — Quartr → SEC exhibit → web ──────────────────────────

def _transcript_from_quartr(ticker: str) -> str | None:
    """Download the most recent earnings call transcript from Quartr API."""
    api_key = os.getenv("QUARTR_API_KEY", "").strip()
    if not api_key:
        return None

    try:
        url = (f"{QUARTR_API_BASE}/documents/transcripts"
               f"?tickers={ticker}&direction=desc&expand=event&limit=10")
        resp = requests.get(url, headers={"x-api-key": api_key}, timeout=20)
        if resp.status_code != 200:
            return None
        data = resp.json()
    except Exception:
        return None

    # Filter to earnings calls
    earnings_type_ids = {26, 27, 28, 29, 35, 36}
    transcripts = [
        d for d in data.get("data", [])
        if d.get("event", {}).get("typeId") in earnings_type_ids
    ]
    if not transcripts:
        return None

    # Dedupe by eventId, prefer typeId=22 (in-house)
    seen = {}
    for d in transcripts:
        eid = d["eventId"]
        if eid not in seen or (d["typeId"] == 22 and seen[eid]["typeId"] != 22):
            seen[eid] = d
    latest = sorted(seen.values(), key=lambda d: d["event"]["date"], reverse=True)[0]

    evt = latest.get("event", {})
    file_url = latest.get("fileUrl")
    if not file_url:
        return None

    try:
        tj = requests.get(file_url, timeout=30).json()
    except Exception:
        return None

    # Build speaker map
    speaker_map = {}
    for entry in tj.get("speaker_mapping", []):
        idx = entry["speaker"]
        sd = entry.get("speaker_data", {})
        name = sd.get("name", f"Speaker {idx}")
        role = sd.get("role", "")
        company = sd.get("company", "")
        if role and company:
            label = f"{name}, {role} ({company})"
        elif role:
            label = f"{name}, {role}"
        else:
            label = name
        speaker_map[idx] = label

    paragraphs = tj.get("transcript", {}).get("paragraphs", [])
    lines = [
        f"# Earnings Call Transcript — {evt.get('title', ticker)} "
        f"({evt.get('date', '')[:10]})",
        f"Source: Quartr\n",
    ]
    last_speaker = None
    for para in paragraphs:
        spk = para.get("speaker")
        text = para.get("text", "").strip()
        if not text:
            continue
        if spk != last_speaker:
            lines.append(f"\n{speaker_map.get(spk, f'Speaker {spk}')}:")
            last_speaker = spk
        lines.append(text)

    return "\n".join(lines)


def _transcript_from_sec_exhibit(ticker: str) -> str | None:
    """Look for a transcript exhibit (ex99.2 or similar) in the most recent 8-K."""
    cik = get_cik(ticker)
    if not cik:
        return None

    try:
        resp = requests.get(
            f"https://data.sec.gov/submissions/CIK{cik}.json",
            headers=EDGAR_HEADERS, timeout=15,
        )
        recent = resp.json().get("filings", {}).get("recent", {})
    except Exception:
        return None

    forms = recent.get("form", [])
    accessions = recent.get("accessionNumber", [])
    dates = recent.get("filingDate", [])

    for i, form in enumerate(forms[:20]):
        if form != "8-K":
            continue
        docs = _get_filing_docs(cik, accessions[i])
        name, url = _find_exhibit(docs, [
            r"ex99.*transcript", r"ex992\b", r"ex-992\b",
            r"transcript", r"exhibit.?99.?2",
        ])
        if not url:
            continue
        raw = _fetch_limited(url, max_bytes=300_000)
        if raw is None:
            continue
        decoded = raw.decode("utf-8", errors="ignore")
        text = _html_to_text(decoded) if "<html" in decoded.lower() else decoded
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return (
            f"# Earnings Call Transcript — {ticker} 8-K {dates[i]}\n"
            f"Source: SEC EDGAR exhibit ({name})\n\n{text}"
        )

    return None


def _transcript_from_audio(ticker: str) -> str | None:
    """Find the earnings call audio/video URL from Quartr events and transcribe via Whisper API.

    Requires QUARTR_API_KEY (to find the recording URL) and OPENAI_API_KEY (for Whisper).
    Downloads up to 25 MB of audio to stay within Whisper's file size limit.
    """
    quartr_key = os.getenv("QUARTR_API_KEY", "").strip()
    openai_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not quartr_key or not openai_key:
        return None

    # Find the most recent earnings event with a recording URL
    try:
        url = (f"{QUARTR_API_BASE}/events"
               f"?tickers={ticker}&direction=desc&limit=10")
        resp = requests.get(url, headers={"x-api-key": quartr_key}, timeout=20)
        if resp.status_code != 200:
            return None
        events = resp.json().get("data", [])
    except Exception:
        return None

    earnings_type_ids = {26, 27, 28, 29, 35, 36}
    audio_url = None
    event_title = ticker
    event_date = ""

    for evt in events:
        if evt.get("typeId") not in earnings_type_ids:
            continue
        # Look for any audio/video recording URL in the event
        for field in ("audioUrl", "recordingUrl", "videoUrl", "mediaUrl"):
            candidate = evt.get(field)
            if candidate:
                audio_url = candidate
                event_title = evt.get("title", ticker)
                event_date = evt.get("date", "")[:10]
                break
        if audio_url:
            break

    if not audio_url:
        return None

    print(f"    Audio found: {audio_url[:80]}... transcribing via Whisper...")

    # Download audio (25 MB cap — Whisper limit is 25 MB)
    try:
        raw = _fetch_limited(audio_url, max_bytes=25 * 1024 * 1024)
        if not raw or len(raw) < 10_000:
            return None
    except Exception:
        return None

    # Determine file extension from URL or default to mp3
    ext = re.search(r"\.(mp3|mp4|m4a|wav|webm|ogg)(\?|$)", audio_url, re.IGNORECASE)
    ext = ext.group(1) if ext else "mp3"

    # Write to a temp file and send to Whisper
    import tempfile
    try:
        from openai import OpenAI
    except ImportError:
        return None

    try:
        client = OpenAI(api_key=openai_key)
        with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
            tmp.write(raw)
            tmp_path = tmp.name

        with open(tmp_path, "rb") as f:
            whisper_resp = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="text",
            )
        os.unlink(tmp_path)

        text = whisper_resp if isinstance(whisper_resp, str) else str(whisper_resp)
        return (
            f"# Earnings Call Transcript — {event_title} ({event_date})\n"
            f"Source: Quartr audio → Whisper transcription\n\n{text}"
        )
    except Exception as e:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        print(f"    Whisper transcription failed: {e}")
        return None


def _transcript_from_web(ticker: str, company_name: str) -> str | None:
    """Search for a transcript via DuckDuckGo."""
    queries = [
        f"{ticker} {company_name} earnings call transcript full text 2025",
        f"{ticker} earnings call transcript Q4 2024 site:seekingalpha.com OR site:fool.com",
    ]
    results = []
    for q in queries:
        r = web_search(q, max_results=4)
        if r and "Search failed" not in r:
            results.append(r)
    if not results:
        return None
    text = "\n\n".join(results)
    return (
        f"# Earnings Call — {ticker} (web search excerpts)\n"
        f"Source: DuckDuckGo\n\n{text}"
    )


def gather_transcript(ticker: str, company_name: str) -> str:
    """Fetch earnings call transcript with automatic fallback chain:
    1. Quartr API (richest — named speakers, full text)
    2. SEC EDGAR 8-K exhibit (some companies file transcripts)
    3. DuckDuckGo web search (summary excerpts)
    """
    print(f"  Fetching transcript: trying Quartr...")
    result = _transcript_from_quartr(ticker)
    if result:
        print(f"    Quartr: OK")
        if len(result) > MAX_CHARS_TRANSCRIPT:
            result = result[:MAX_CHARS_TRANSCRIPT] + "\n\n[... truncated ...]"
        return result

    print(f"    Quartr: unavailable — trying SEC EDGAR transcript exhibit...")
    result = _transcript_from_sec_exhibit(ticker)
    if result:
        print(f"    SEC exhibit: OK")
        if len(result) > MAX_CHARS_TRANSCRIPT:
            result = result[:MAX_CHARS_TRANSCRIPT] + "\n\n[... truncated ...]"
        return result

    print(f"    SEC exhibit: not found — trying Quartr audio + Whisper transcription...")
    result = _transcript_from_audio(ticker)
    if result:
        print(f"    Whisper: OK")
        if len(result) > MAX_CHARS_TRANSCRIPT:
            result = result[:MAX_CHARS_TRANSCRIPT] + "\n\n[... truncated ...]"
        return result

    print(f"    Whisper: unavailable — falling back to web search...")
    result = _transcript_from_web(ticker, company_name)
    if result:
        print(f"    Web search: OK")
        if len(result) > MAX_CHARS_TRANSCRIPT:
            result = result[:MAX_CHARS_TRANSCRIPT] + "\n\n[... truncated ...]"
        return result

    return f"No transcript found for {ticker} (Quartr unavailable, no SEC exhibit, no audio, web search empty)."


# ── File parsing ───────────────────────────────────────────────────────────────

def parse_pdf(file_path: str) -> str:
    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        full_text = "\n\n".join(text_parts).strip()
        if len(full_text) > MAX_CHARS_FILE:
            full_text = full_text[:MAX_CHARS_FILE] + "\n\n[... truncated ...]"
        return full_text
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def parse_excel(file_path: str) -> str:
    try:
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path)
            result = df.to_string(max_rows=100)
        else:
            xl = pd.ExcelFile(file_path)
            parts = []
            for sheet in xl.sheet_names[:5]:
                df = xl.parse(sheet)
                parts.append(f"### Sheet: {sheet}\n{df.to_string(max_rows=80, max_cols=20)}")
            result = "\n\n".join(parts)
        if len(result) > MAX_CHARS_FILE:
            result = result[:MAX_CHARS_FILE] + "\n\n[... truncated ...]"
        return result
    except Exception as e:
        return f"[Excel/CSV extraction failed: {e}]"


def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                text_parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        full_text = "\n\n".join(text_parts).strip()
        if len(full_text) > MAX_CHARS_FILE:
            full_text = full_text[:MAX_CHARS_FILE] + "\n\n[... truncated ...]"
        return full_text
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def parse_uploaded_files(file_paths: list, captions: dict = None) -> str:
    if not file_paths:
        return ""
    captions = captions or {}
    parts = []
    for fp in file_paths:
        if not os.path.exists(fp):
            continue
        ext = Path(fp).suffix.lower()
        name = Path(fp).name
        if ext == ".pdf":
            content = parse_pdf(fp)
        elif ext in (".docx", ".doc"):
            content = parse_docx(fp)
        elif ext in (".xlsx", ".xls", ".csv"):
            content = parse_excel(fp)
        else:
            content = f"[Unsupported file type: {ext}]"
        header = f"## Uploaded File: {name}"
        if fp in captions and captions[fp]:
            header += f"\nUser context: {captions[fp]}"
        parts.append(f"{header}\n{content}")
    return "\n\n".join(parts)


# ── Telegram digest search ──────────────────────────────────────────────────────

TELEGRAM_API_ID   = 33919151
TELEGRAM_API_HASH = "dd0a935bd6545cf56910292ff4445c4e"
MAX_CHARS_TELEGRAM = 12000


async def _fetch_telegram_messages(group_name: str, search_query: str, days: int, max_messages: int) -> str:
    from telethon import TelegramClient
    from telethon.sessions import StringSession
    from datetime import datetime, timedelta, timezone

    session_str = os.environ.get("TELEGRAM_SESSION", "")
    if not session_str:
        return "TELEGRAM_SESSION not set."

    session = StringSession(session_str) if len(session_str) > 20 else session_str
    tg = TelegramClient(session, TELEGRAM_API_ID, TELEGRAM_API_HASH)

    try:
        await tg.connect()
        if not await tg.is_user_authorized():
            return "Telegram session not authorized."

        dialogs = await tg.get_dialogs()
        group_entity = None
        for d in dialogs:
            if d.name == group_name:
                group_entity = d.entity
                break

        if group_entity is None:
            return f"Group '{group_name}' not found."

        cutoff = datetime.now(timezone.utc) - timedelta(days=days)
        messages = []

        async for msg in tg.iter_messages(group_entity, search=search_query, limit=max_messages):
            if msg.date < cutoff:
                break
            if msg.text and msg.text.strip():
                messages.append(f"[{msg.date.strftime('%Y-%m-%d')}] {msg.text.strip()}")

        if not messages:
            return f"No messages found for '{search_query}' in '{group_name}' (last {days} days)."

        result = f"## {group_name} — '{search_query}' ({len(messages)} messages, last {days} days)\n\n"
        result += "\n\n---\n\n".join(messages)
        if len(result) > MAX_CHARS_TELEGRAM:
            result = result[:MAX_CHARS_TELEGRAM] + "\n\n[... truncated ...]"
        return result

    finally:
        await tg.disconnect()


def gather_telegram_digest(group_name: str, search_query: str, days: int, max_messages: int = 200) -> str:
    """Fetch relevant messages from a Telegram digest group. Safe to call from a thread executor."""
    import asyncio
    try:
        return asyncio.run(_fetch_telegram_messages(group_name, search_query, days, max_messages))
    except Exception as e:
        return f"Telegram digest fetch failed: {e}"


# ── Main entry point ───────────────────────────────────────────────────────────

def gather_all(ticker: str, staged_files: list, staged_captions: dict = None) -> dict:
    """
    Gather all data for a ticker. Returns a dict with keys:
    yfinance, news, sec, transcript, uploaded_files, telegram_digest, company_name
    """
    print(f"  Gathering yfinance data for {ticker}...")
    yf_data = gather_yfinance(ticker)

    company_name = ticker
    try:
        info = yf.Ticker(ticker).info
        company_name = info.get("longName", ticker)
    except Exception:
        pass

    print(f"  Searching news for {company_name}...")
    news_data = gather_news(ticker, company_name)

    print(f"  Downloading SEC EDGAR filings (8-K press release, 10-Q, 10-K)...")
    sec_data = gather_sec(ticker)

    print(f"  Fetching earnings transcript...")
    transcript_data = gather_transcript(ticker, company_name)

    print(f"  Searching Stock Digest (Telegram) for {ticker} / {company_name} — last 12 months...")
    tg_ticker = gather_telegram_digest("📈 Stock Digest", ticker, days=365, max_messages=200)
    tg_name   = gather_telegram_digest("📈 Stock Digest", company_name, days=365, max_messages=100) \
                if company_name != ticker else ""
    telegram_data = tg_ticker
    if tg_name and "No messages found" not in tg_name:
        telegram_data += "\n\n" + tg_name

    print(f"  Parsing {len(staged_files)} staged file(s)...")
    uploaded = parse_uploaded_files(staged_files, staged_captions)

    return {
        "ticker": ticker,
        "company_name": company_name,
        "yfinance": yf_data,
        "news": news_data,
        "sec": sec_data,
        "transcript": transcript_data,
        "telegram_digest": telegram_data,
        "uploaded_files": uploaded,
    }
